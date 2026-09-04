import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

def shade(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{margin}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)

doc = Document()

# Set page margins
section = doc.sections[0]
section.top_margin = Cm(1.6)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# Configure Typography Styles
styles = doc.styles
normal_style = styles["Normal"]
normal_style.font.name = "Segoe UI"
normal_style.font.size = Pt(11)
normal_style.paragraph_format.space_after = Pt(6)
normal_style.paragraph_format.line_spacing = 1.18

title_style = styles["Title"]
title_style.font.name = "Segoe UI"
title_style.font.size = Pt(22)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(27, 43, 107) # #1b2b6b

h1_style = styles["Heading 1"]
h1_style.font.name = "Segoe UI"
h1_style.font.size = Pt(14)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(30, 58, 138) # #1e3a8a

h2_style = styles["Heading 2"]
h2_style.font.name = "Segoe UI"
h2_style.font.size = Pt(12)
h2_style.font.bold = True
h2_style.font.color.rgb = RGBColor(14, 116, 144) # #0e7490

# Title Block
title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("KỊCH BẢN THUYẾT TRÌNH 3 PHÚT (180 GIÂY)")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run("Luồng Hoạt Động Tuyển Dụng Thông Minh (iRSA) — Bám Sát Slide Trực Quan")
sub_run.bold = True
sub_run.font.size = Pt(12)
sub_run.font.color.rgb = RGBColor(59, 130, 246)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta.add_run("⏱ Tổng thời lượng: 03:00 · 🎯 Tốc độ: ~140 từ/phút · 📊 10 Slide chuẩn hóa")
meta_run.font.size = Pt(10)
meta_run.font.color.rgb = RGBColor(100, 116, 139)

# Summary Box Table
table = doc.add_table(rows=1, cols=1)
table.alignment = WD_ALIGN_PARAGRAPH.CENTER
cell = table.cell(0, 0)
set_cell_margins(cell, 120, 150, 120, 150)
p_box = cell.paragraphs[0]
p_box.paragraph_format.space_after = Pt(2)
run_box_b = p_box.add_run("📌 NGUYÊN TẮC THUYẾT TRÌNH 3 PHÚT:\n")
run_box_b.bold = True
run_box_b.font.color.rgb = RGBColor(30, 58, 138)
p_box.add_run(
    "1. Dõng dạc, đúng nhịp: Phần LỜI THOẠI được canh chuẩn xác cho 180 giây; không nói thừa chữ.\n"
    "2. Thao tác dứt khoát: Nhấn phím 'Phím Cách' hoặc 'Mũi tên phải' ngay khi hết mỗi ý để khớp visual trên màn hình.\n"
    "3. Điểm đắt giá cần nhấn: (1) Công thức từ khóa minh bạch; (2) Lọc cơ bản giúp tiết kiệm 70% chi phí & tăng tốc 12x; (3) Khâu phỏng vấn: Thu âm ➔ Bóc script ➔ Chấm rubric."
)
shade(p_box, "F0F9FF")

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# 8 Main Script Sections for 10 Slides (180 seconds)
SCRIPT_SECTIONS = [
    {
        "index": "01",
        "time": "00:00 – 00:15 (15 giây)",
        "slide": "Slide 1: Giới thiệu quy trình tuyển dụng thông minh",
        "action": "Giữ Slide 1 trên màn hình. Giọng tự tin, chào hội đồng và nêu ngay giá trị cốt lõi.",
        "speech": (
            "Kính chào Ban giám khảo và mọi người. Em xin đại diện nhóm trình bày về Luồng hoạt động "
            "của Hệ thống Tuyển dụng Thông minh iRSA. Đây là giải pháp kết hợp giữa công thức toán học "
            "minh bạch và AI chuyên sâu, giúp tự động hóa toàn diện từ lúc nhận CV cho đến khi HR ra "
            "quyết định tuyển dụng cuối cùng."
        ),
        "emphasis": "iRSA: Chu trình tuyển dụng tự động, minh bạch và có căn cứ dữ liệu.",
        "words": 65
    },
    {
        "index": "02",
        "time": "00:15 – 00:35 (20 giây)",
        "slide": "Slide 2: Thực trạng & Điểm đau (Pain Points)",
        "action": "Chuyển sang Slide 2. Chỉ tay vào 3 con số thống kê đỏ rực trên màn hình.",
        "speech": (
            "Hiện nay, quy trình tuyển dụng truyền thống đang gặp bế tắc lớn: 75% CV bị các bộ lọc cũ "
            "loại oan vì chỉ bắt đúng chữ mà không hiểu ngữ cảnh; HR mất từ 15 đến 20 giờ mỗi tuần "
            "đọc CV thủ công nhưng vẫn mắc bẫy ứng viên nhồi từ khóa ảo. Trong khi đó, hơn 80% ứng viên "
            "rơi vào hố đen tuyển dụng, nộp hồ sơ mà không hề nhận được bất kỳ phản hồi nào."
        ),
        "emphasis": "Điểm nghẽn kép: HR kiệt sức vì đọc CV rác — Ứng viên bức xúc vì thiếu minh bạch.",
        "words": 82
    },
    {
        "index": "03",
        "time": "00:35 – 00:55 (20 giây)",
        "slide": "Slide 3: Sơ đồ luồng hoạt động 6 bước khép kín",
        "action": "Chuyển sang Slide 3. Quét nhanh từ bước 1 đến bước 6 trên pipeline.",
        "speech": (
            "Để giải quyết triệt để vấn đề trên, iRSA xây dựng luồng 6 bước khép kín: Ứng viên nộp CV; "
            "Nhà tuyển dụng chấm so khớp từ khóa có công thức; Ứng viên tự quyết định có chấm AI hay không; "
            "Đạt 2 lượt chấm sẽ được tự động xếp lịch; Sau đó vào phòng phỏng vấn trực tuyến với 3 khâu "
            "thu âm, bóc script và chấm điểm; Cuối cùng, HR nhìn vào bảng điểm tổng hợp để đưa ra quyết định."
        ),
        "emphasis": "Quy trình 6 chặng liên hoàn, dữ liệu thông suốt từ CV đến kết quả phỏng vấn.",
        "words": 85
    },
    {
        "index": "04",
        "time": "00:55 – 01:25 (30 giây)",
        "slide": "Slide 4: Bước 1 & 2 — Nộp CV & Công thức so khớp từ khóa chi tiết",
        "action": "Chuyển sang Slide 4. Chỉ vào công thức toán học và kéo thử nhẹ thanh trượt Live Calculator.",
        "speech": (
            "Điểm đặc biệt đầu tiên là khâu So khớp từ khóa được lượng hóa bằng công thức toán học cố định, "
            "hoàn toàn giải thích được: Điểm kỹ năng chiếm 60% — trong đó kỹ năng Must-have có trọng số "
            "gấp đôi Nice-to-have; Kinh nghiệm chiếm 30% tính theo số năm thực tế; và Học vấn chiếm 10%. "
            "Nhờ bổ sung từ điển từ đồng nghĩa và embedding vector, hệ thống không bỏ sót nhân tài chỉ vì "
            "cách viết khác nhau."
        ),
        "emphasis": "Không còn là 'hộp đen bí mật' — Điểm sàng lọc có công thức toán học minh bạch 100%.",
        "words": 88
    },
    {
        "index": "05",
        "time": "01:25 – 01:55 (30 giây)",
        "slide": "Slide 5 & 6: Tại sao cần lọc cơ bản trước? Bài toán chi phí & Hiệu năng",
        "action": "Chuyển Slide 5 rồi lướt sang Slide 6. Nhấn mạnh vào bảng so sánh chi phí và độ trễ 50ms.",
        "speech": (
            "Tại sao không đưa thẳng 100% CV vào LLM? Bởi vì 70% CV nộp vào là trái ngành hoặc thiếu chuẩn. "
            "Nếu cho tất cả vào AI, với 1.000 CV ta mất gần 50 USD và nghẽn API đến 2 tiếng. "
            "Bằng việc lọc cơ bản trước trên CPU nội bộ với chi phí 0 đồng, hệ thống loại bỏ ngay 70% hồ sơ "
            "không đạt trong 5 giây, tiết kiệm 70% ngân sách API, giảm độ trễ từ 8 giây xuống dưới 50 mili-giây, "
            "và loại bỏ hoàn toàn nguy cơ sập hệ thống do lỗi Rate Limit."
        ),
        "emphasis": "Lọc cơ bản là tấm khiên kinh tế và kỹ thuật: Tiết kiệm 70% chi phí, nhanh hơn 12 lần.",
        "words": 98
    },
    {
        "index": "06",
        "time": "01:55 – 02:15 (20 giây)",
        "slide": "Slide 7: Bước 3 — Quyết định chấm AI của Ứng viên (Opt-in)",
        "action": "Chuyển sang Slide 7. Nhấn vào hai nhánh lựa chọn A và B.",
        "speech": (
            "Sau khi biết điểm từ khóa, quyền chủ động thuộc về Ứng viên: Ứng viên có thể chọn dừng lại "
            "hoặc kích hoạt AI Agent đánh giá sâu. Khi kích hoạt, AI sẽ bóc tách dẫn chứng dự án thực tế "
            "thay vì đếm từ khóa, vạch trần các trường hợp nhồi nhét từ khóa ảo và tự động sinh sẵn bộ câu "
            "hỏi phỏng vấn dựa đúng vào những thiếu hụt của hồ sơ đó."
        ),
        "emphasis": "Trao quyền cho ứng viên — AI kiểm tra bằng chứng thực chiến, chống gian lận.",
        "words": 82
    },
    {
        "index": "07",
        "time": "02:15 – 02:40 (25 giây)",
        "slide": "Slide 8 & 9: Bước 4 & 5 — Hẹn lịch tự động & Phòng phỏng vấn (3 Khâu)",
        "action": "Lướt Slide 8 sang Slide 9. Chỉ vào 3 khâu: Thu âm, Bóc script và Chấm điểm trên giao diện phỏng vấn.",
        "speech": (
            "Chỉ khi đạt cả 2 lượt chấm, ứng viên mới được hệ thống tự động đồng bộ lịch và gửi link phòng họp. "
            "Tại phòng phỏng vấn trực tuyến, hệ thống thực hiện 3 khâu chuẩn hóa: Thứ nhất, Thu âm trực tiếp câu "
            "trả lời qua mic; Thứ hai, Bóc script tự động bằng Whisper sang văn bản chính xác từng câu; và "
            "Thứ ba, AI phân tích script để chấm điểm theo phương pháp STAR và thang rubric cố định."
        ),
        "emphasis": "Chuẩn hóa phỏng vấn: Thu âm ➔ Bóc script ➔ Chấm điểm rubric khách quan.",
        "words": 88
    },
    {
        "index": "08",
        "time": "02:40 – 03:00 (20 giây)",
        "slide": "Slide 10: Bước 6 — Chấm điểm cuối cùng & HR ra quyết định",
        "action": "Chuyển sang Slide 10. Chỉ vào nút Offer / Reject và bảng điểm tổng hợp.",
        "speech": (
            "Cuối cùng, hệ thống tổng hợp điểm số theo tỷ trọng: 25% Từ khóa, 35% Đánh giá AI và 40% Phỏng vấn. "
            "HR Dashboard trực quan hóa toàn bộ năng lực và đề xuất gợi ý. Nhưng quan trọng nhất: "
            "theo nguyên tắc Human-in-the-loop, AI chỉ đóng vai trò trợ lý cung cấp bằng chứng, "
            "chính HR mới là người nhấn nút Offer hay Reject cuối cùng. Em xin trân trọng cảm ơn!"
        ),
        "emphasis": "Human-in-the-loop: AI tham mưu bằng chứng — Con người ra quyết định tuyển dụng.",
        "words": 80
    }
]

# Write Sections
for item in SCRIPT_SECTIONS:
    h = doc.add_heading(f"Phần {item['index']}: {item['slide']}", level=1)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(3)

    # Meta timeline
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(4)
    r_time = p_meta.add_run(f"⏱ Thời lượng: {item['time']}  |  📝 Độ dài: ~{item['words']} từ")
    r_time.bold = True
    r_time.font.size = Pt(10)
    r_time.font.color.rgb = RGBColor(234, 88, 12) # orange

    # Action box
    p_act = doc.add_paragraph()
    p_act.paragraph_format.left_indent = Cm(0.3)
    p_act.paragraph_format.right_indent = Cm(0.3)
    p_act.paragraph_format.space_after = Pt(6)
    p_act_tag = p_act.add_run("🎬 THAO TÁC TRÌNH CHIẾU: ")
    p_act_tag.bold = True
    p_act_tag.font.size = Pt(9.5)
    p_act_tag.font.color.rgb = RGBColor(71, 85, 105)
    r_act = p_act.add_run(item["action"])
    r_act.font.size = Pt(9.5)
    shade(p_act, "F1F5F9")

    # Speech box
    p_speech_label = doc.add_paragraph()
    p_speech_label.paragraph_format.space_after = Pt(2)
    r_speech_lbl = p_speech_label.add_run("🗣 LỜI THOẠI DIỄN THUYẾT:")
    r_speech_lbl.bold = True
    r_speech_lbl.font.color.rgb = RGBColor(37, 99, 235)

    p_speech = doc.add_paragraph()
    p_speech.paragraph_format.first_line_indent = Cm(0.5)
    p_speech.paragraph_format.space_after = Pt(6)
    p_speech.paragraph_format.line_spacing = 1.22
    r_speech = p_speech.add_run(f'"{item["speech"]}"')
    r_speech.font.size = Pt(11)

    # Emphasis box
    p_emp = doc.add_paragraph()
    p_emp.paragraph_format.left_indent = Cm(0.3)
    p_emp.paragraph_format.right_indent = Cm(0.3)
    p_emp.paragraph_format.space_after = Pt(10)
    p_emp_tag = p_emp.add_run("⭐ CÂU CHỐT / ĐIỂM NHẤN: ")
    p_emp_tag.bold = True
    p_emp_tag.font.size = Pt(9.5)
    p_emp_tag.font.color.rgb = RGBColor(16, 185, 129)
    r_emp = p_emp.add_run(item["emphasis"])
    r_emp.bold = True
    r_emp.font.size = Pt(9.5)
    shade(p_emp, "ECFDF5")

# Add Appendix Page for Q&A
doc.add_section(WD_SECTION.NEW_PAGE)
h_app = doc.add_heading("Phụ Lục: Cẩm Nang Trả Lời Phản Biện Của Hội Đồng (Q&A)", level=1)
h_app.paragraph_format.space_before = Pt(10)
h_app.paragraph_format.space_after = Pt(8)

p_intro_qa = doc.add_paragraph()
p_intro_qa.add_run("Dưới đây là 4 câu hỏi giám khảo thường chất vấn nhất đối với luồng tuyển dụng này và câu trả lời chuẩn xác:")

QA_LIST = [
    (
        "Câu hỏi 1: Tại sao không dùng AI chấm điểm ngay từ đầu mà phải qua bước lọc từ khóa bằng công thức?",
        "Trả lời: Nếu gửi 100% CV vào LLM, hệ thống sẽ gặp 2 rủi ro chí mạng: (1) Chi phí token bùng nổ vì 70% CV nộp vào là spam hoặc thiếu hẳn kỹ năng bắt buộc; (2) Nghẽn hàng đợi và dính lỗi Rate Limit 429 khi có đợt nộp dồn dập. Bằng cách dùng bộ lọc có công thức cố định (Rule-based), hệ thống loại ngay 70% hồ sơ không đạt trong <50ms với chi phí 0 đồng, bảo vệ ngân sách và giữ hàng đợi AI luôn thông thoáng."
    ),
    (
        "Câu hỏi 2: Bộ lọc từ khóa có nguy cơ loại bỏ nhầm ứng viên giỏi do khác biệt cách diễn đạt không?",
        "Trả lời: Không loại nhầm nhờ cơ chế Hybrid Matching: (1) Từ điển từ đồng nghĩa đa tầng (Skill Aliases: JS ↔ JavaScript, K8s ↔ Kubernetes); (2) Cosine Similarity trên Semantic Vector Embeddings (ngưỡng tương đồng >= 0.82) để bắt trọn ngữ cảnh dù ứng viên dùng từ khác; (3) Phân tách rõ Must-have (nhân đôi trọng số) và Nice-to-have, đảm bảo đánh giá toàn diện."
    ),
    (
        "Câu hỏi 3: Tại sao lại để ứng viên quyết định có chấm AI hay không (Candidate Opt-in)?",
        "Trả lời: Tính năng này mang lại lợi ích 3 bên: (1) Ứng viên nắm quyền chủ động và minh bạch thông tin sau khi biết điểm từ khóa; (2) Tiết kiệm chi phí vận hành (những ứng viên thấy điểm quá thấp hoặc không còn nhu cầu sẽ tự rút, hệ thống không tốn credit gọi LLM); (3) Tối ưu tỷ lệ chuyển đổi khi chỉ tập trung vào nhóm ứng viên thực sự quan tâm và nghiêm túc."
    ),
    (
        "Câu hỏi 4: Cơ chế phỏng vấn: 'Thu âm -> Bóc script -> Chấm điểm' hoạt động như thế nào về mặt kỹ thuật?",
        "Trả lời: Quy trình gồm 3 chặng khép kín: (1) Client-side Audio Stream thu âm qua WebRTC, khử ồn môi trường và lọc khoảng lặng (VAD); (2) Speech-to-Text Pipeline sử dụng mô hình Whisper bóc băng âm thanh thành text có timestamp với độ chính xác >95%; (3) LLM Rubric Engine phân tích văn bản theo cấu trúc STAR (Situation, Task, Action, Result) và rubric kỹ năng để cho điểm khách quan."
    ),
    (
        "Câu hỏi 5: Làm thế nào để đảm bảo tính công bằng và loại bỏ thiên vị (bias) khi AI chấm điểm phỏng vấn?",
        "Trả lời: (1) Chấm trên script văn bản thuần túy, loại bỏ các yếu tố ngoại hình, giọng điệu vùng miền, giới tính; (2) Rubric chấm điểm được chuẩn hóa cố định với tiêu chí rõ ràng (thang điểm 1-5 theo STAR); (3) System Prompt được thiết kế chặt chẽ chống ảo giác (hallucination); (4) Tính năng Human-in-the-loop: HR có thể nghe lại audio gốc và xem giải trình của AI để kiểm chứng."
    ),
    (
        "Câu hỏi 6: Nếu ứng viên cố tình gian lận (đọc tài liệu có sẵn, nhờ người khác trả lời) thì sao?",
        "Trả lời: Hệ thống có cơ chế kiểm soát: (1) Giới hạn thời gian chuẩn bị và trả lời (Time-boxed) cho từng câu hỏi; (2) Tạo câu hỏi phỏng vấn động cá nhân hóa dựa riêng trên từng dự án trong CV, khiến ứng viên không thể học thuộc vẹt hay chép tài liệu tổng quát; (3) Phân tích độ trễ phản xạ và tính liền mạch của luồng âm thanh."
    ),
    (
        "Câu hỏi 7: Công thức tính điểm tổng kết có thể tùy chỉnh theo từng vị trí tuyển dụng không?",
        "Trả lời: Có, công thức S_Final = w1*S_KW + w2*S_AI + w3*S_Interview cho phép HR cấu hình linh hoạt theo Job Family: Ví dụ đối với vị trí Sales/CSKH cần giao tiếp mạnh, tỷ trọng Phỏng vấn có thể nâng lên 55%; đối với vị trí Backend Engineer chuyên sâu, điểm AI phân tích CV kỹ thuật có thể đẩy lên 45%."
    ),
    (
        "Câu hỏi 8: Tại sao HR vẫn là người ra quyết định cuối cùng thay vì để hệ thống tự động hoàn toàn?",
        "Trả lời: Đây là triết lý AI-Assisted, không phải AI-Replaced. Tuyển dụng gắn liền với văn hóa doanh nghiệp và trách nhiệm pháp lý. AI đóng vai trò như một trợ lý lọc nhiễu, chuẩn bị dữ liệu và khuyến nghị khoa học, còn quyết định Offer/Reject thuộc thẩm quyền của HR Manager để đảm bảo tính nhân văn và phù hợp tổ chức."
    ),
    (
        "Câu hỏi 9: Dự tính chi phí và khả năng mở rộng (Scalability) của hệ thống khi có quy mô lớn?",
        "Trả lời: Nhờ kiến trúc 2 tầng (Two-stage Filtering), tầng 1 dùng thuật toán nhẹ chạy trên CPU xử lý hàng ngàn CV trong vài giây với chi phí gần bằng 0. Chỉ ~30% CV chất lượng mới vào tầng LLM. Với 10,000 CV/tháng, doanh nghiệp chỉ tốn ~$135 chi phí API thay vì $450+, tiết kiệm hơn 70% ngân sách và đảm bảo throughput ổn định."
    ),
    (
        "Câu hỏi 10: Hệ thống lưu trữ dữ liệu âm thanh và CV của ứng viên ra sao để tuân thủ quyền riêng tư?",
        "Trả lời: (1) Dữ liệu CV và file ghi âm được mã hóa tại chỗ (AES-256) và truyền qua TLS 1.3; (2) Chính sách lưu trữ dữ liệu có thời hạn (Data Retention Policy: tự động xóa file ghi âm sau 30-90 ngày theo cấu hình doanh nghiệp); (3) Quyền truy cập phân tầng (RBAC) nghiêm ngặt, tuân thủ Nghị định 13/2023/NĐ-CP về bảo vệ dữ liệu cá nhân."
    )
]

for q, a in QA_LIST:
    p_q = doc.add_paragraph()
    p_q.paragraph_format.space_before = Pt(8)
    p_q.paragraph_format.space_after = Pt(2)
    r_q = p_q.add_run(q)
    r_q.bold = True
    r_q.font.color.rgb = RGBColor(30, 58, 138)

    p_a = doc.add_paragraph()
    p_a.paragraph_format.first_line_indent = Cm(0.4)
    p_a.paragraph_format.space_after = Pt(6)
    p_a.add_run(a)
    shade(p_a, "F8FAFC")

# Save document
# Save document safely handling file locks if opened in MS Word
paths_to_save = [
    Path("D:/VIN/team/iRSA/Kich-ban-thuyet-trinh-3-phut-iRSA.docx"),
    Path("D:/VIN/team/iRSA/presentation/Kich-ban-thuyet-trinh-3-phut-iRSA.docx"),
    Path("C:/Users/Admin/.gemini/antigravity-cli/brain/dcfd27b0-7fd4-48d5-a7e0-def491deed55/Kich-ban-thuyet-trinh-3-phut-iRSA.docx")
]

for p in paths_to_save:
    try:
        doc.save(p)
        print(f"Successfully saved to: {p}")
    except PermissionError:
        # If open in MS Word, save with alternate backup name
        alt_p = p.parent / f"{p.stem}_updated{p.suffix}"
        try:
            doc.save(alt_p)
            print(f"File {p.name} is currently open in Word. Saved updated version to: {alt_p}")
        except Exception as e:
            print(f"Failed to save {alt_p}: {e}")
    except Exception as e:
        print(f"Failed to save {p}: {e}")
