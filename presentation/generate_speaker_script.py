from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(__file__).with_name("Kich-ban-thuyet-trinh-iRSA-5-phut.docx")


SECTIONS = [
    {
        "time": "00:00–00:15",
        "title": "Mở đầu — Giới thiệu iRSA",
        "action": "Hiển thị slide 1. Giữ nguyên slide trong khi chào và giới thiệu.",
        "speech": (
            "Xin chào mọi người, nhóm chúng em xin giới thiệu iRSA – Intelligent Resume Screening "
            "Automation. Đây là nền tảng tuyển dụng tích hợp AI, giúp biến mỗi hồ sơ thành một quyết "
            "định có căn cứ."
        ),
        "emphasis": "Tuyển đúng người — có bằng chứng.",
    },
    {
        "time": "00:15–00:40",
        "title": "Pain point của bài toán",
        "action": "Chuyển slide 2, lần lượt chỉ vào hai điểm nghẽn: thời gian và tính nhất quán.",
        "speech": (
            "Bài toán tập trung vào hai điểm nghẽn. Thứ nhất, với mỗi hồ sơ, HR phải lặp lại việc "
            "đọc và đối chiếu với JD nên thời gian xử lý tăng theo số ứng viên. Thứ hai, cùng một CV "
            "có thể nhận kết luận khác nhau khi thiếu rubric và bằng chứng chung. Vì vậy, nhu cầu "
            "không chỉ là xử lý nhanh hơn mà còn phải chuẩn hóa cách đánh giá."
        ),
        "emphasis": "Hai điểm nghẽn cốt lõi: thời gian xử lý và tính nhất quán của đánh giá.",
    },
    {
        "time": "00:40–01:00",
        "title": "Mục tiêu hướng đến của dự án",
        "action": "Chuyển slide 3 và chỉ vào ba mục tiêu chính.",
        "speech": (
            "Mục tiêu ban đầu của iRSA là giảm tải và chuẩn hóa khâu chấm CV. Tuy nhiên, một điểm "
            "số vẫn cần được xác minh. Vì vậy nhóm mở rộng sản phẩm thành quy trình khép kín: từ hồ "
            "sơ, bằng chứng, chấm điểm đến câu hỏi, phỏng vấn và đánh giá câu trả lời. AI hỗ trợ "
            "quyết định; HR vẫn là người đưa ra quyết định cuối cùng."
        ),
        "emphasis": "Từ CV đến quyết định phỏng vấn trong một quy trình khép kín.",
    },
    {
        "time": "01:00–01:40",
        "title": "Bước 1 — Tiếp nhận và chuẩn hóa CV",
        "action": (
            "Mở http://localhost:4300 → chọn công việc → xem JD → Ứng tuyển → chọn CV → "
            "xác nhận nộp hồ sơ."
        ),
        "speech": (
            "Đầu tiên, ứng viên chọn một vị trí, xem yêu cầu và nộp CV có sẵn hoặc tải CV mới. "
            "Hệ thống tiếp nhận hồ sơ, trích xuất nội dung và liên kết với tiêu chí của JD. Sau đó "
            "iRSA tạo Quick Matching để HR có tín hiệu phù hợp ban đầu. Ở bước này hệ thống chưa "
            "gọi LLM, vì vậy mọi CV đều có thể được sàng lọc với chi phí thấp."
        ),
        "emphasis": "CV được chuẩn hóa và Quick Matching trước khi sử dụng AI chuyên sâu.",
    },
    {
        "time": "01:40–02:10",
        "title": "Bước 2 — Cổng chấm điểm hai tầng",
        "action": (
            "Chuyển sang http://localhost:4200 → Applications → xem điểm matching → bấm Chấm "
            "điểm AI hoặc mở kết quả đã chuẩn bị sẵn."
        ),
        "speech": (
            "Ở phía HR, điểm Quick Matching giúp sắp xếp và ưu tiên hồ sơ mà chưa cần dùng LLM. "
            "Nếu CV nào cũng tự động chạy AI thì chi phí token sẽ tăng đáng kể. Vì vậy HR chỉ bấm "
            "Chấm điểm AI khi cần xem phân tích sâu, bằng chứng và khuyến nghị. Trong demo, em có thể "
            "bấm chấm điểm hoặc mở kết quả đã chuẩn bị sẵn để không phải chờ dịch vụ AI."
        ),
        "emphasis": "HR quyết định hồ sơ nào cần sử dụng chi phí AI chuyên sâu.",
    },
    {
        "time": "02:10–03:00",
        "title": "Bước 3 — Pipeline chấm điểm AI",
        "action": "Sau khi bấm Chấm điểm AI, chuyển slide 6 và mô tả bốn bước của pipeline.",
        "speech": (
            "Khi HR bấm Chấm điểm AI, pipeline chạy qua bốn bước. Extractor chia CV theo cấu trúc "
            "và giữ vị trí từng đoạn. Evidence Retrieval dùng vector search tìm nội dung liên quan "
            "cho mỗi yêu cầu. Evaluator áp dụng rubric để đánh giá kỹ năng, kinh nghiệm và học vấn. "
            "Cuối cùng, Verifier kiểm tra bằng chứng có thực sự nằm trong CV hay không, đồng thời đối "
            "chiếu điểm với khuyến nghị. Nếu kết quả chưa nhất quán, hệ thống yêu cầu đánh giá lại."
        ),
        "emphasis": "Extractor → Evidence Retrieval → Evaluator → Verifier.",
    },
    {
        "time": "03:00–03:35",
        "title": "Bước 4 — Kết quả có giải thích",
        "action": "Chuyển slide 7 → chỉ vào công thức → hai kỹ năng → hai bằng chứng.",
        "speech": (
            "Kết quả không chỉ là một con số. Điểm được tính theo rubric: kỹ năng bắt buộc chiếm "
            "50 phần trăm, kỹ năng ưu tiên 20 phần trăm, kinh nghiệm 20 phần trăm và học vấn 10 phần "
            "trăm. Với từng kỹ năng, HR thấy confidence và đoạn bằng chứng tương ứng. Không có bằng "
            "chứng thì confidence bằng 0; thiếu must-have thì điểm bị giới hạn. Nhờ đó HR hiểu vì sao "
            "hồ sơ phù hợp hoặc cần xác minh thêm."
        ),
        "emphasis": "Confidence là độ tin cậy của bằng chứng, không phải điểm phù hợp.",
    },
    {
        "time": "03:35–04:15",
        "title": "Bước 5 — Quy trình phỏng vấn khép kín",
        "action": "Chuyển slide 8; đi từ Sinh câu hỏi → HR duyệt → Ghi âm → STT → Đánh giá.",
        "speech": (
            "Sau khi hoàn thiện luồng chấm CV, nhóm nhận thấy một điểm số vẫn cần được xác minh trong "
            "thực tế. Vì vậy iRSA được mở rộng thành quy trình khép kín. Hệ thống biến điểm mạnh và "
            "khoảng trống trong CV thành câu hỏi cá nhân hóa. HR duyệt câu hỏi và lập lịch. Interview "
            "room ghi âm câu trả lời, Speech-to-Text tạo transcript, sau đó AI đánh giá từng câu trả "
            "lời và tổng hợp toàn bộ buổi phỏng vấn. Dữ liệu screening nhờ đó được sử dụng tiếp, thay "
            "vì dừng lại ở một điểm số."
        ),
        "emphasis": "Screening tạo giả thuyết; phỏng vấn giúp HR xác minh giả thuyết đó.",
    },
    {
        "time": "04:15–04:40",
        "title": "Kỹ thuật, kiến trúc và triển khai",
        "action": "Chuyển slide 9. Trình bày nhanh từ trái sang phải, không đi sâu từng công nghệ.",
        "speech": (
            "Về kiến trúc, Candidate Portal và Admin Portal được tách theo nhóm người dùng. FastAPI "
            "điều phối API, scoring pipeline và interview workflow. PostgreSQL lưu dữ liệu nghiệp vụ, "
            "object storage lưu CV và audio, Redis hỗ trợ cache và rate limiting. Hệ thống có Pydantic "
            "validation, JWT HttpOnly cookie, Swagger, test suite và tài liệu kiến trúc."
        ),
        "emphasis": "Kiến trúc hỗ trợ xuyên suốt từ CV, scoring đến interview.",
    },
    {
        "time": "04:40–05:00",
        "title": "Kết luận",
        "action": "Chuyển slide 10; nhìn khán giả, không tiếp tục thao tác trên hệ thống.",
        "speech": (
            "Tóm lại, iRSA giải quyết hai điểm nghẽn là thời gian sàng lọc và sự thiếu nhất quán. "
            "Sản phẩm dùng Quick Matching để tối ưu chi phí, AI Agent để chấm điểm có bằng chứng, "
            "sau đó nối kết quả screening với phỏng vấn để tạo thành quy trình tuyển dụng khép kín. "
            "HR luôn giữ quyền kiểm tra và đưa ra quyết định cuối cùng."
        ),
        "emphasis": "Từ CV đến quyết định — nhanh hơn, có căn cứ và khép kín.",
    },
]


def shade(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{margin}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.7)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(11.5)
styles["Normal"].paragraph_format.space_after = Pt(7)
styles["Normal"].paragraph_format.line_spacing = 1.18
styles["Title"].font.name = "Aptos Display"
styles["Title"].font.size = Pt(30)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = RGBColor(31, 58, 170)
styles["Heading 1"].font.name = "Aptos Display"
styles["Heading 1"].font.size = Pt(19)
styles["Heading 1"].font.bold = True
styles["Heading 1"].font.color.rgb = RGBColor(21, 38, 73)
styles["Heading 2"].font.name = "Aptos Display"
styles["Heading 2"].font.size = Pt(15)
styles["Heading 2"].font.bold = True
styles["Heading 2"].font.color.rgb = RGBColor(52, 86, 241)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("KỊCH BẢN THUYẾT TRÌNH DEMO iRSA")
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Thời lượng mục tiêu: 5 phút · 10 slide · Có thao tác demo")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(96, 112, 138)

intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
intro_run = intro.add_run("Mẹo luyện tập: đọc phần LỜI THOẠI, còn phần THAO TÁC DEMO chỉ là lời nhắc hành động.")
intro_run.italic = True
intro_run.font.color.rgb = RGBColor(71, 85, 105)

for index, item in enumerate(SECTIONS, start=1):
    heading = doc.add_heading(f"{index}. {item['title']}", level=1)
    heading.paragraph_format.space_before = Pt(12)
    time_p = doc.add_paragraph()
    time_p.paragraph_format.space_after = Pt(6)
    time_run = time_p.add_run(f"⏱  {item['time']}")
    time_run.bold = True
    time_run.font.color.rgb = RGBColor(255, 105, 70)

    action = doc.add_paragraph()
    action.paragraph_format.left_indent = Cm(.35)
    action.paragraph_format.right_indent = Cm(.35)
    action.paragraph_format.space_before = Pt(2)
    action.paragraph_format.space_after = Pt(9)
    action.add_run("THAO TÁC DEMO  ").bold = True
    action.add_run(item["action"])
    shade(action, "EEF1FF")

    label = doc.add_paragraph()
    label.paragraph_format.space_after = Pt(3)
    label_run = label.add_run("LỜI THOẠI")
    label_run.bold = True
    label_run.font.color.rgb = RGBColor(52, 86, 241)

    speech = doc.add_paragraph(item["speech"])
    speech.paragraph_format.first_line_indent = Cm(.6)
    speech.paragraph_format.line_spacing = 1.25
    speech.paragraph_format.space_after = Pt(7)

    emphasis = doc.add_paragraph()
    emphasis.paragraph_format.left_indent = Cm(.35)
    emphasis.paragraph_format.right_indent = Cm(.35)
    emphasis.add_run("CÂU NHẤN: ").bold = True
    emp_run = emphasis.add_run(item["emphasis"])
    emp_run.bold = True
    shade(emphasis, "F1F8DF")

doc.add_section(WD_SECTION.NEW_PAGE)
doc.add_heading("Phụ lục — Trả lời nhanh khi giám khảo hỏi", level=1)

qa = [
    (
        "Công thức điểm phù hợp là gì?",
        "Score = 0.5M + 0.2N + 0.2E + 0.1D. M là tỷ lệ must-have, N là nice-to-have, "
        "E là mức đáp ứng kinh nghiệm và D là học vấn; các thành phần đều chuẩn hóa về thang 0–100.",
    ),
    (
        "Điểm khác confidence như thế nào?",
        "Điểm thể hiện mức phù hợp với công việc. Confidence thể hiện độ tin cậy của bằng chứng mà "
        "hệ thống tìm được. Không có bằng chứng thì confidence bằng 0.",
    ),
    (
        "AI có tự quyết định tuyển hay không?",
        "Không. iRSA hỗ trợ sàng lọc và giải thích kết quả; HR vẫn kiểm tra bằng chứng, phỏng vấn và "
        "đưa ra quyết định cuối cùng.",
    ),
    (
        "Hệ thống hạn chế hallucination thế nào?",
        "Mỗi nhận định kỹ năng phải có trích dẫn từ CV. Verifier kiểm tra grounding, tính nhất quán "
        "giữa điểm và khuyến nghị; kết quả chưa đạt sẽ được yêu cầu đánh giá lại.",
    ),
    (
        "Nếu dịch vụ AI lỗi trong lúc demo?",
        "Quick Matching vẫn cho HR tín hiệu sàng lọc ban đầu. Với AI Evaluation, nhóm đã chuẩn bị một "
        "application có kết quả sẵn để tiếp tục trình bày mà không phải chờ LLM.",
    ),
]

for question, answer in qa:
    p = doc.add_paragraph()
    q = p.add_run(question)
    q.bold = True
    q.font.color.rgb = RGBColor(31, 58, 170)
    doc.add_paragraph(answer)

doc.add_heading("Checklist 2 phút trước khi trình bày", level=1)
checklist = [
    "Mở sẵn Candidate Portal và Admin Portal ở hai tab.",
    "Đăng nhập sẵn tài khoản Candidate và HR.",
    "Chuẩn bị CV mẫu và application đã có kết quả AI.",
    "Chọn trước hai bằng chứng kỹ năng và hai câu hỏi nổi bật.",
    "Mở slide ở chế độ toàn màn hình; nhấn T để chạy đồng hồ.",
    "Tắt thông báo và đóng các tab không liên quan.",
]
for item in checklist:
    doc.add_paragraph(item, style="List Bullet")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run("iRSA · Intelligent Resume Screening Automation · Kịch bản 5 phút")
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(120, 132, 151)

doc.save(OUTPUT)
print(f"Created: {OUTPUT.resolve()}")
